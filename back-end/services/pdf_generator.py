from __future__ import annotations

from typing import Union, Tuple
from io import BytesIO
from fpdf import FPDF

try:
    from docx import Document  # optional
    _DOCX_AVAILABLE = True
except Exception:
    _DOCX_AVAILABLE = False


def _to_latin1(text: str) -> str:
    # FPDF core font supports latin-1; replace unsupported glyphs
    return text.encode("latin-1", "replace").decode("latin-1")


def generate_resume_pdf(resume_text: str) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, _to_latin1("Resume"), ln=True)

    pdf.set_font("Arial", size=11)
    # Use multi_cell to wrap long text
    safe_text = _to_latin1(resume_text or "(No content provided)")
    for paragraph in safe_text.split("\n\n"):
        pdf.multi_cell(0, 6, paragraph)
        pdf.ln(1)

    # Output as bytes
    pdf_bytes: Union[bytes, bytearray, str] = pdf.output(dest="S").encode("latin-1")
    if isinstance(pdf_bytes, str):
        pdf_bytes = pdf_bytes.encode("latin-1", "replace")
    return bytes(pdf_bytes)


def generate_resume_bytes(resume_text: str, fmt: str) -> Tuple[bytes, str, str]:
    fmt = (fmt or "pdf").lower()
    if fmt == "pdf":
        pdf_bytes = generate_resume_pdf(resume_text)
        return pdf_bytes, "application/pdf", "resume.pdf"
    if fmt == "txt":
        return (resume_text or "").encode("utf-8"), "text/plain; charset=utf-8", "resume.txt"
    if fmt == "docx":
        if not _DOCX_AVAILABLE:
            raise RuntimeError("DOCX export requires python-docx. Please install it.")
        doc = Document()
        doc.add_heading("Resume", level=1)
        for para in (resume_text or "").split("\n\n"):
            doc.add_paragraph(para)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "resume.docx"
    raise ValueError("Unsupported format")
